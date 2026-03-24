import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_markdown_to_docx(doc, md_text):
    # Normalize line endings
    md_text = md_text.replace('\r\n', '\n')
    lines = md_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Handle empty lines
        if not line:
            i += 1
            continue
            
        # Handle Tables
        # A table starts if the line contains '|' and the next line is a separator line
        if '|' in line and i + 1 < len(lines):
            next_line = lines[i+1].strip()
            # Separator line: |---|---| or ---|--- or | :--- | ---: |
            if re.match(r'^\|?[\s\-\|:]+\|?$', next_line) and '|' in next_line:
                # We found a table
                table_rows = []
                # Collect header
                table_rows.append(line)
                # Collect separator
                table_rows.append(next_line)
                
                i += 2
                # Collect data rows
                while i < len(lines) and '|' in lines[i]:
                    table_rows.append(lines[i].strip())
                    i += 1
                
                _render_table(doc, table_rows)
                continue

        # Handle Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        # Handle Lists
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
            p = doc.add_paragraph(style='List Bullet')
            _parse_inline_formatting(p, line[2:].strip())
        elif re.match(r'^\d+[\.\)]\s', line):
            p = doc.add_paragraph(style='List Number')
            # Remove both 1. and 1) formats
            clean_line = re.sub(r'^\d+[\.\)]\s+', '', line).strip()
            _parse_inline_formatting(p, clean_line)
        # Handle Bold/Italic lines that are not part of anything else
        else:
            p = doc.add_paragraph()
            _parse_inline_formatting(p, line)
        
        i += 1

def _render_table(doc, table_rows):
    if not table_rows:
        return
    
    # Parse columns for all rows first to determine max columns
    parsed_rows = []
    for row in table_rows:
        # Skip separator rows in the final table data
        if re.match(r'^\|?[\s\-\|:]+\|?$', row) and '|' in row:
            continue
            
        content = row.strip()
        if content.startswith('|'):
            content = content[1:]
        if content.endswith('|'):
            content = content[:-1]
        
        # Split by '|' but avoid splitting on escaped pipes if any (though rare in this context)
        cols = [col.strip() for col in content.split('|')]
        if any(c for c in cols): # Only add if there's some content
            parsed_rows.append(cols)
        
    if not parsed_rows:
        return
        
    num_cols = max(len(row) for row in parsed_rows)
    if num_cols == 0:
        return

    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(parsed_rows):
        for j, col_text in enumerate(row):
            if j < num_cols:
                try:
                    cell = table.cell(i, j)
                    # Clear existing text
                    cell.text = ''
                    p = cell.paragraphs[0]
                    _parse_inline_formatting(p, col_text)
                    if i == 0:
                        # Header row
                        for run in p.runs:
                            run.bold = True
                except Exception:
                    pass

def _parse_inline_formatting(paragraph, text):
    # regex to find ***bold_italic***, **bold**, *italic*
    # Tokenize by bold/italic markers
    pattern = r'(\*\*\*.*?\*\*\*|___.*?___|\*\*.*?\*\*|__.*?__|(?<!\*)\*.*?\*(?!\*)|\b_.*?_\b)'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token:
            continue
        
        if (token.startswith('***') and token.endswith('***')) or (token.startswith('___') and token.endswith('___')):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif (token.startswith('**') and token.endswith('**')) or (token.startswith('__') and token.endswith('__')):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif (token.startswith('*') and token.endswith('*')) or (token.startswith('_') and token.endswith('_')):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)
