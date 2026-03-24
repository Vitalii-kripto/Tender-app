import os
import logging
import zipfile
from typing import List, Dict, Any
from .document_service import DocumentService
from .legal_analysis_service import LegalAnalysisService
from backend.config import DOCUMENTS_ROOT
from .job_service import job_service

logger = logging.getLogger("BatchAnalysis")

def analyze_tenders_batch_job(
    job_id: str, 
    tender_ids: List[str], 
    doc_service: DocumentService, 
    legal_service: LegalAnalysisService, 
    selected_files: Dict[str, List[str]] = None
):
    """
    Основной воркер для пакетного анализа тендеров.
    Реализует Word-only архитектуру:
    1. Извлечение текста (DocumentService).
    2. Полнотекстовый ИИ-анализ (LegalAnalysisService).
    3. Генерация Word-отчета.
    4. Сохранение результатов в JobService.
    5. Создание ZIP-архива для пакета.
    """
    selected_files = selected_files or {}
    documents_root = DOCUMENTS_ROOT
    report_paths = []
    
    for tid in tender_ids:
        logger.info(f"--- [START TENDER ANALYSIS: {tid}] ---")
        tender_dir = os.path.join(documents_root, tid)
        
        # 0. Инициализация статуса
        job_service.update_tender_stage(job_id, tid, "Подготовка документов", 10)
        
        # 1. Проверка выбора файлов
        requested_files = selected_files.get(tid, [])
        if not requested_files:
            logger.warning(f"No files selected for tender {tid}")
            job_service.complete_tender(job_id, tid, {
                "status": "error",
                "final_report_markdown": "Ошибка: не выбрано ни одного файла для анализа. Пожалуйста, выберите хотя бы один документ.",
                "summary_notes": "Файлы не выбраны.",
                "file_statuses": [],
                "export_available": False
            })
            continue

        # 2. Проверка существования директории
        if not os.path.exists(tender_dir):
            logger.warning(f"Tender directory not found: {tender_dir}")
            job_service.complete_tender(job_id, tid, {
                "status": "error",
                "final_report_markdown": "Ошибка: директория с документами не найдена. Возможно, тендер еще не был обработан или файлы были удалены.",
                "summary_notes": "Директория не найдена.",
                "file_statuses": [{"filename": f, "status": "error", "message": "Директория не найдена"} for f in requested_files],
                "export_available": False
            })
            continue
            
        # 3. Извлечение текста
        job_service.update_tender_stage(job_id, tid, "Извлечение текста", 20)
        
        files_data = []
        file_statuses = []
        available_files = os.listdir(tender_dir)
        
        for filename in requested_files:
            if filename not in available_files:
                logger.warning(f"File {filename} not found in {tender_dir}")
                file_statuses.append({"filename": filename, "status": "error", "message": "Файл не найден на диске"})
                continue
                
            filepath = os.path.join(tender_dir, filename)
            try:
                text = doc_service.extract_text(filepath)
                if text and len(text.strip()) > 0:
                    files_data.append({"filename": filename, "text": text})
                    file_statuses.append({"filename": filename, "status": "ok", "message": "Текст успешно извлечен"})
                else:
                    file_statuses.append({"filename": filename, "status": "warning", "message": "Файл пуст или текст не извлечен"})
            except Exception as e:
                logger.error(f"Failed to extract text from {filename}: {e}")
                file_statuses.append({"filename": filename, "status": "error", "message": f"Ошибка извлечения: {str(e)}"})

        if not files_data:
            logger.error(f"No text extracted from any of the selected files for tender {tid}")
            job_service.complete_tender(job_id, tid, {
                "status": "error",
                "final_report_markdown": "Ошибка: не удалось извлечь текст ни из одного выбранного файла. Проверьте форматы документов.",
                "summary_notes": "Текст не извлечен.",
                "file_statuses": file_statuses,
                "export_available": False
            })
            continue

        # 4. ИИ-анализ (LegalAnalysisService)
        try:
            def stage_callback(stage, progress, status="running"):
                job_service.update_tender_stage(job_id, tid, stage, progress, status)

            analysis_result = legal_service.analyze_tender(
                files_data, 
                tender_id=tid, 
                callback=stage_callback
            )
            
            final_markdown = analysis_result.get('final_report_markdown', '')
            summary_notes = analysis_result.get('summary_notes', '')
            cleaned_context_len = analysis_result.get('cleaned_context_len', 0)
            final_report_len = analysis_result.get('final_report_len', 0)
            
            # 5. Генерация Word-отчета
            report_path = "N/A"
            export_available = False
            try:
                from docx import Document
                from backend.markdown_parser import add_markdown_to_docx
                from docx.shared import Pt
                
                doc = Document()
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(11)
                
                doc.add_heading(f'Юридическое заключение по тендеру {tid}', 0)
                
                if final_markdown:
                    add_markdown_to_docx(doc, final_markdown)
                else:
                    doc.add_paragraph("Отчет пуст.")
                
                os.makedirs(tender_dir, exist_ok=True)
                report_filename = f"report_{tid}.docx"
                report_path = os.path.abspath(os.path.join(tender_dir, report_filename))
                doc.save(report_path)
                report_paths.append(report_path)
                export_available = True
            except Exception as e:
                logger.error(f"Error generating Word report for tender {tid}: {e}")

            # 6. Финальное логирование требуемых метрик
            logger.info(f"--- [ANALYSIS LOGS FOR TENDER {tid}] ---")
            logger.info(f"- Cleaned context length: {cleaned_context_len} chars")
            logger.info(f"- Final markdown report length: {final_report_len} chars")
            logger.info(f"- Word report path: {report_path}")
            logger.info(f"----------------------------------------")

            # 7. Завершение задачи для тендера
            job_service.complete_tender(job_id, tid, {
                "status": analysis_result.get('status', 'success'),
                "final_report_markdown": final_markdown,
                "summary_notes": summary_notes,
                "file_statuses": file_statuses,
                "report_path": report_path,
                "export_available": export_available
            })
            logger.info(f"--- [END TENDER ANALYSIS: {tid}] ---")
            
        except Exception as e:
            logger.error(f"Analysis failed for tender {tid}: {e}", exc_info=True)
            job_service.complete_tender(job_id, tid, {
                "status": "error",
                "final_report_markdown": f"Критическая ошибка анализа: {str(e)}",
                "summary_notes": "Ошибка анализа.",
                "file_statuses": file_statuses,
                "export_available": False
            })
            
    # 8. Создание ZIP-архива для пакета (если больше 1 тендера)
    zip_path = "N/A"
    if len(report_paths) > 1:
        try:
            batch_dir = os.path.join(documents_root, "batch_results")
            os.makedirs(batch_dir, exist_ok=True)
            zip_filename = f"batch_{job_id}.zip"
            zip_path = os.path.abspath(os.path.join(batch_dir, zip_filename))
            
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for r_path in report_paths:
                    if os.path.exists(r_path):
                        zipf.write(r_path, os.path.basename(r_path))
            
            logger.info(f"--- [BATCH ZIP LOG] ---")
            logger.info(f"- Batch ZIP path: {zip_path}")
            logger.info(f"------------------------")
        except Exception as e:
            logger.error(f"Error creating batch ZIP: {e}")

    # 9. Проверка завершения всего задания
    job_service.check_job_completion(job_id)
