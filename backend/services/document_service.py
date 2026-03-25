import os
import re
import io
import numpy as np
from pypdf import PdfReader
from fastapi import UploadFile
import aiofiles
import platform
import logging

from backend.logger import logger

import pypdfium2 as pdfium
from paddleocr import PaddleOCR

class DocumentService:
    """
    –°–µ—А–≤–Є—Б –і–ї—П —А–∞–±–Њ—В—Л —Б –і–Њ–Ї—Г–Љ–µ–љ—В–∞–Љ–Є.
    –Я–Њ–і–і–µ—А–ґ–Є–≤–∞–µ—В –Є–Ј–≤–ї–µ—З–µ–љ–Є–µ —В–µ–Ї—Б—В–∞ –Є–Ј PDF –Є OCR (—А–∞—Б–њ–Њ–Ј–љ–∞–≤–∞–љ–Є–µ —Б–Ї–∞–љ–Њ–≤) —З–µ—А–µ–Ј PaddleOCR.
    """
    UPLOAD_DIR = "uploaded_docs"

    def __init__(self):
        os.makedirs(self.UPLOAD_DIR, exist_ok=True)
        try:
            # –Ш–љ–Є—Ж–Є–∞–ї–Є–Ј–∞—Ж–Є—П PaddleOCR (–Љ–Њ–і–µ–ї–Є —Б–Ї–∞—З–Є–≤–∞—О—В—Б—П –њ—А–Є –њ–µ—А–≤–Њ–Љ –Ј–∞–њ—Г—Б–Ї–µ)
            self.ocr_engine = PaddleOCR(use_angle_cls=True, lang='ru', show_log=False)
            logger.info("PaddleOCR engine initialized successfully.")
        except Exception as e:
            logger.error(f"Failed to initialize PaddleOCR: {e}")
            raise RuntimeError(f"OCR Configuration Error: Failed to initialize PaddleOCR: {e}")

    async def save_file(self, file: UploadFile) -> str:
        """–°–Њ—Е—А–∞–љ—П–µ—В –Ј–∞–≥—А—Г–ґ–µ–љ–љ—Л–є —Д–∞–є–ї"""
        file_path = os.path.join(self.UPLOAD_DIR, file.filename)
        logger.info(f"Saving file to: {file_path}")
        async with aiofiles.open(file_path, 'wb') as out_file:
            content = await file.read()
            await out_file.write(content)
        return file_path

    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """
        –Я—Л—В–∞–µ—В—Б—П –Ї–Њ–љ–≤–µ—А—В–Є—А–Њ–≤–∞—В—М .doc –≤ .docx.
        –Т–Њ–Ј–≤—А–∞—Й–∞–µ—В –њ—Г—В—М –Ї –љ–Њ–≤–Њ–Љ—Г —Д–∞–є–ї—Г –Є–ї–Є –Њ—А–Є–≥–Є–љ–∞–ї—М–љ—Л–є –њ—Г—В—М, –µ—Б–ї–Є –љ–µ —Г–і–∞–ї–Њ—Б—М.
        """
        docx_path = doc_path + "x"
        if os.path.exists(docx_path):
            logger.info(f"DOCX version already exists: {docx_path}")
            return docx_path

        # 1. –Я–Њ–њ—Л—В–Ї–∞ —З–µ—А–µ–Ј win32com (—В–Њ–ї—М–Ї–Њ Windows + Word)
        if platform.system() == "Windows":
            try:
                import win32com.client as win32
                import pythoncom
                # –Ш–љ–Є—Ж–Є–∞–ї–Є–Ј–∞—Ж–Є—П COM –і–ї—П —В–µ–Ї—Г—Й–µ–≥–Њ –њ–Њ—В–Њ–Ї–∞ (–≤–∞–ґ–љ–Њ –і–ї—П FastAPI/async)
                pythoncom.CoInitialize()
                
                word = win32.gencache.EnsureDispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(doc_path))
                # 16 = wdFormatXMLDocument (.docx)
                doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)
                doc.Close()
                word.Quit()
                logger.info(f"Successfully converted {doc_path} to {docx_path} using MS Word.")
                return docx_path
            except Exception as e:
                logger.warning(f"win32com conversion failed (check if Word is installed): {e}")
            finally:
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass

        # 2. –Я–Њ–њ—Л—В–Ї–∞ —З–µ—А–µ–Ј –Є–Ј–≤–ї–µ—З–µ–љ–Є–µ —В–µ–Ї—Б—В–∞ –Є —Б–Њ–Ј–і–∞–љ–Є–µ –љ–Њ–≤–Њ–≥–Њ .docx (fallback)
        text = self._extract_text_from_doc(doc_path)
        if text and not text.startswith("[–Ю–®–Ш–С–Ъ–Р"):
            try:
                import docx
                new_doc = docx.Document()
                new_doc.add_paragraph(f"--- –Р–Т–Ґ–Ю–Ь–Р–Ґ–Ш–І–Х–°–Ъ–Р–ѓ –Ъ–Ю–Э–Т–Х–†–Ґ–Р–¶–Ш–ѓ –Ш–Ч .DOC ---\n–Ю—А–Є–≥–Є–љ–∞–ї: {os.path.basename(doc_path)}\n\n")
                new_doc.add_paragraph(text)
                new_doc.save(docx_path)
                logger.info(f"Created {docx_path} from extracted text of {doc_path}")
                return docx_path
            except Exception as e:
                logger.error(f"Failed to create docx from text: {e}")

        return doc_path

    def extract_text(self, file_path: str) -> str:
        """
        –£–Љ–љ–Њ–µ –Є–Ј–≤–ї–µ—З–µ–љ–Є–µ —В–µ–Ї—Б—В–∞ —Б –Ј–∞—Й–Є—В–Њ–є –Њ—В —Б–±–Њ–µ–≤.
        –Я–Њ–і–і–µ—А–ґ–Є–≤–∞–µ—В PDF, DOCX, XLSX, XLS, DOC.
        """
        ext = os.path.splitext(file_path)[1].lower()
        logger.info(f"--- [START EXTRACTION] ---")
        logger.info(f"File path: {file_path}")
        logger.info(f"Extension: {ext}")
        
        full_text = ""
        handler = "None"

        try:
            # 1. –Ю–±—А–∞–±–Њ—В–Ї–∞ .doc (—Б –њ–Њ–њ—Л—В–Ї–Њ–є –Ї–Њ–љ–≤–µ—А—В–∞—Ж–Є–Є –Є –і–µ–і—Г–њ–ї–Є–Ї–∞—Ж–Є–Є)
            if ext == '.doc':
                docx_path = file_path + "x"
                if os.path.exists(docx_path):
                    logger.info(f"Skipping .doc extraction because .docx already exists: {docx_path}")
                    file_path = docx_path
                    ext = '.docx'
                else:
                    logger.info(f"Handler: .doc legacy converter/antiword")
                    handler = "doc_converter"
                    new_path = self._convert_doc_to_docx(file_path)
                    if new_path.endswith('.docx'):
                        logger.info(f"Successfully converted .doc to .docx: {new_path}")
                        file_path = new_path
                        ext = '.docx'
                        # –Я—А–Њ–і–Њ–ї–ґ–∞–µ–Љ –Ї–∞–Ї .docx –љ–Є–ґ–µ
                    else:
                        logger.info(f"Conversion failed or not applicable, using fallback extraction for .doc")
                        full_text = self._extract_text_from_doc(file_path)
                        logger.info(f"Extraction complete. Handler: {handler}. Characters: {len(full_text)}")
                        return full_text

            # 2. –Ю–±—А–∞–±–Њ—В–Ї–∞ .docx
            if ext == '.docx':
                handler = "python-docx"
                import docx
                doc = docx.Document(file_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # 3. –Ю–±—А–∞–±–Њ—В–Ї–∞ .xlsx
            elif ext == '.xlsx':
                handler = "openpyxl"
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                sheets_text = []
                for sheet in wb.worksheets:
                    sheet_data = []
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip().replace("|", "").strip():
                            sheet_data.append(row_text)
                    if sheet_data:
                        sheets_text.append(f"=== –Ы–Ш–°–Ґ: {sheet.title} ===\n" + "\n".join(sheet_data))
                full_text = "\n\n".join(sheets_text)
            
            # 4. –Ю–±—А–∞–±–Њ—В–Ї–∞ .xls
            elif ext == '.xls':
                handler = "xlrd"
                import xlrd
                wb = xlrd.open_workbook(file_path)
                sheets_text = []
                for sheet in wb.sheets():
                    sheet_data = []
                    for row_idx in range(sheet.nrows):
                        row = sheet.row_values(row_idx)
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip().replace("|", "").strip():
                            sheet_data.append(row_text)
                    if sheet_data:
                        sheets_text.append(f"=== –Ы–Ш–°–Ґ: {sheet.name} ===\n" + "\n".join(sheet_data))
                full_text = "\n\n".join(sheets_text)
            
            # 5. –Ю–±—А–∞–±–Њ—В–Ї–∞ .pdf
            elif ext == '.pdf':
                handler = "pypdf"
                reader = PdfReader(file_path)
                text_pages = []
                for page in reader.pages:
                    try:
                        extracted = page.extract_text()
                        if extracted:
                            text_pages.append(extracted)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from a PDF page: {e}")
                
                full_text = "\n".join(text_pages)
                
                # –Ю—Ж–µ–љ–Ї–∞ –Ї–∞—З–µ—Б—В–≤–∞ –Є–Ј–≤–ї–µ—З–µ–љ–љ–Њ–≥–Њ —В–µ–Ї—Б—В–∞
                is_quality_good = self._is_text_quality_good(full_text)
                char_count = len(full_text)
                
                logger.info(f"--- [PDF SPECIFIC LOGS] ---")
                logger.info(f"Text Layer Quality: {'GOOD' if is_quality_good else 'POOR'}")
                logger.info(f"OCR Triggered: {'Yes' if not is_quality_good else 'No'}")
                
                if not is_quality_good:
                    logger.info(f"PDF text quality is POOR (Chars: {char_count}). Triggering OCR fallback...")
                    ocr_result = self._perform_ocr(file_path)
                    
                    logger.info(f"OCR fallback successful. Text obtained after OCR: {len(ocr_result)} chars")
                    handler = "pypdf + pypdfium2 + PaddleOCR (OCR)"
                    full_text = ocr_result
                else:
                    logger.info(f"PDF text quality is GOOD (Chars: {char_count}). Using native text layer.")
                logger.info(f"---------------------------")

            # 6. –Э–µ–њ–Њ–і–і–µ—А–ґ–Є–≤–∞–µ–Љ—Л–є —Д–Њ—А–Љ–∞—В
            else:
                logger.warning(f"Unsupported file format: {ext} for file {file_path}")
                return f"[SYSTEM INFO] –§–Њ—А–Љ–∞—В {ext} –љ–µ –њ–Њ–і–і–µ—А–ґ–Є–≤–∞–µ—В—Б—П."

            char_count = len(full_text)
            ocr_used = handler.endswith("(OCR)")
            
            logger.info(f"--- [EXTRACTION SUMMARY] ---")
            logger.info(f"File Name: {os.path.basename(file_path)}")
            logger.info(f"File Type: {ext}")
            logger.info(f"Handler Used: {handler}")
            logger.info(f"Characters Extracted: {char_count}")
            logger.info(f"OCR Used: {'Yes' if ocr_used else 'No'}")
            logger.info(f"----------------------------")
            
            return full_text

        except Exception as e:
            logger.error(f"Extraction failed for {file_path} using {handler}. Error: {str(e)}", exc_info=True)
            return f"[ERROR] –Ю—И–Є–±–Ї–∞ –њ—А–Є —З—В–µ–љ–Є–Є {ext} ({handler}): {str(e)}"

    def _is_text_quality_good(self, text: str) -> bool:
        """
        –Я—А–Њ–≤–µ—А—П–µ—В –Ї–∞—З–µ—Б—В–≤–Њ –Є–Ј–≤–ї–µ—З–µ–љ–љ–Њ–≥–Њ —В–µ–Ї—Б—В–∞.
        –Т–Њ–Ј–≤—А–∞—Й–∞–µ—В True, –µ—Б–ї–Є —В–µ–Ї—Б—В –Ї–∞—З–µ—Б—В–≤–µ–љ–љ—Л–є, –Є False, –µ—Б–ї–Є —В—А–µ–±—Г–µ—В—Б—П OCR.
        –£–ґ–µ—Б—В–Њ—З–µ–љ–љ—Л–µ –Ї—А–Є—В–µ—А–Є–Є –і–ї—П PDF.
        """
        if not text or len(text.strip()) < 150: # –£–≤–µ–ї–Є—З–Є–ї–Є –њ–Њ—А–Њ–≥ –Љ–Є–љ–Є–Љ–∞–ї—М–љ–Њ–є –і–ї–Є–љ—Л
            return False

        total_chars = len(text)
        
        # 1. –Ф–Њ–ї—П –Љ—Г—Б–Њ—А–љ—Л—Е —Б–Є–Љ–≤–Њ–ї–Њ–≤
        clean_pattern = r'[a-zA-Z–∞-—П–Р-–ѓ—С–Б0-9\s\.,!?;:()""\'\'\-\+=\[\]/\\<>@#\$%\^&\*¬Ђ¬ївДЦ]'
        clean_chars_count = len(re.findall(clean_pattern, text))
        garbage_ratio = 1 - (clean_chars_count / total_chars) if total_chars > 0 else 1
        
        # 2. –Ф–Њ–ї—П –љ–Њ—А–Љ–∞–ї—М–љ—Л—Е —Б–ї–Њ–≤ –љ–∞ —А—Г—Б—Б–Ї–Њ–Љ
        words = text.split()
        if not words:
            return False
            
        russian_words = [w for w in words if re.search(r'[–∞-—П–Р-–ѓ—С–Б]{3,}', w)]
        russian_words_ratio = len(russian_words) / len(words) if words else 0
        
        # 3. –Э–∞–ї–Є—З–Є–µ –і–ї–Є–љ–љ—Л—Е –Є—Б–њ–Њ—А—З–µ–љ–љ—Л—Е —Б—В—А–Њ–Ї
        max_word_len = max(len(w) for w in words)
        avg_word_len = sum(len(w) for w in words) / len(words)
        
        # 4. –Ф–Њ–ї—П –љ–µ—З–Є—В–∞–µ–Љ—Л—Е —Д—А–∞–≥–Љ–µ–љ—В–Њ–≤
        unreadable_chars = len(re.findall(r'[\?\x00-\x08\x0b\x0c\x0e-\x1f]', text))
        unreadable_ratio = unreadable_chars / total_chars if total_chars > 0 else 0

        logger.info(f"PDF Quality Metrics: Garbage={garbage_ratio:.2f}, RusWords={russian_words_ratio:.2f}, MaxWord={max_word_len}, AvgWord={avg_word_len:.1f}, Unreadable={unreadable_ratio:.2f}")

        # –£–ґ–µ—Б—В–Њ—З–µ–љ–љ—Л–µ –њ–Њ—А–Њ–≥–Њ–≤—Л–µ –Ј–љ–∞—З–µ–љ–Є—П:
        # - –Ь—Г—Б–Њ—А–∞ > 10% (–±—Л–ї–Њ 15%)
        # - –†—Г—Б—Б–Ї–Є—Е —Б–ї–Њ–≤ < 30% (–±—Л–ї–Њ 20%)
        # - –°–ї–Є—И–Ї–Њ–Љ –і–ї–Є–љ–љ—Л–µ "—Б–ї–Њ–≤–∞" (> 80 —Б–Є–Љ–≤–Њ–ї–Њ–≤, –±—Л–ї–Њ 100)
        # - –°–ї–Є—И–Ї–Њ–Љ –Љ–љ–Њ–≥–Њ –љ–µ—З–Є—В–∞–µ–Љ—Л—Е –Ј–љ–∞–Ї–Њ–≤ (> 3%, –±—Л–ї–Њ 5%)
        
        if garbage_ratio > 0.10:
            return False
        if russian_words_ratio < 0.30:
            # –Х—Б–ї–Є —Н—В–Њ –љ–µ —В–µ—Е–љ–Є—З–µ—Б–Ї–Є–є –і–Њ–Ї—Г–Љ–µ–љ—В –љ–∞ –∞–љ–≥–ї–Є–є—Б–Ї–Њ–Љ
            english_words = [w for w in words if re.search(r'[a-zA-Z]{3,}', w)]
            english_words_ratio = len(english_words) / len(words)
            if english_words_ratio < 0.40: # –Ш –љ–µ –∞–љ–≥–ї–Є–є—Б–Ї–Є–є —В–Њ–ґ–µ
                return False
        if max_word_len > 80 or avg_word_len > 18:
            return False
        if unreadable_ratio > 0.03:
            return False

        return True

    def _perform_ocr(self, file_path: str) -> str:
        """–Т—Б–њ–Њ–Љ–Њ–≥–∞—В–µ–ї—М–љ—Л–є –Љ–µ—В–Њ–і –і–ї—П OCR —А–∞—Б–њ–Њ–Ј–љ–∞–≤–∞–љ–Є—П —З–µ—А–µ–Ј pypdfium2 –Є PaddleOCR"""
        logger.info(f"Starting OCR for {file_path} using pypdfium2 + PaddleOCR")
        
        # –Ю—В–Ї—А—Л–≤–∞–µ–Љ PDF —З–µ—А–µ–Ј pypdfium2
        pdf = pdfium.PdfDocument(file_path)
        num_pages = len(pdf)
        # –Ю–≥—А–∞–љ–Є—З–Є–≤–∞–µ–Љ –Ї–Њ–ї–Є—З–µ—Б—В–≤–Њ —Б—В—А–∞–љ–Є—Ж –і–ї—П OCR (–њ–µ—А–≤—Л–µ 20 –і–ї—П –±–∞–ї–∞–љ—Б–∞ —Б–Ї–Њ—А–Њ—Б—В–Є –Є –Ї–∞—З–µ—Б—В–≤–∞)
        pages_to_process = min(num_pages, 20)
        
        ocr_text_pages = []
        
        for i in range(pages_to_process):
            logger.info(f"Processing page {i+1}/{pages_to_process}...")
            page = pdf[i]
            
            # –†–µ–љ–і–µ—А–Є–Љ —Б—В—А–∞–љ–Є—Ж—Г –≤ –Є–Ј–Њ–±—А–∞–ґ–µ–љ–Є–µ (scale=2 –і–ї—П 144 DPI, scale=3 –і–ї—П 216 DPI)
            # PaddleOCR —Е–Њ—А–Њ—И–Њ —А–∞–±–Њ—В–∞–µ—В –љ–∞ 144-200 DPI
            bitmap = page.render(scale=2)
            pil_image = bitmap.to_pil()
            
            # –Ъ–Њ–љ–≤–µ—А—В–Є—А—Г–µ–Љ PIL Image –≤ numpy array –і–ї—П PaddleOCR
            img_array = np.array(pil_image)
            
            # –Т—Л–њ–Њ–ї–љ—П–µ–Љ OCR
            result = self.ocr_engine.ocr(img_array, cls=True)
            
            page_text = []
            if result and result[0]:
                for line in result[0]:
                    # line[1][0] - —Н—В–Њ —В–µ–Ї—Б—В, line[1][1] - —Н—В–Њ —Г–≤–µ—А–µ–љ–љ–Њ—Б—В—М (confidence)
                    text_line = line[1][0]
                    page_text.append(text_line)
            
            ocr_text_pages.append(f"--- Page {i+1} ---\n" + "\n".join(page_text))
        
        pdf.close()
        
        if ocr_text_pages:
            full_ocr_text = "\n\n".join(ocr_text_pages)
            logger.info(f"OCR successful. Extracted {len(full_ocr_text)} characters from {pages_to_process} pages.")
            return full_ocr_text
        else:
            logger.warning("OCR ran but found no text.")
            return "[OCR WARNING] OCR –Њ—В—А–∞–±–Њ—В–∞–ї, –љ–Њ —В–µ–Ї—Б—В –љ–µ –љ–∞–є–і–µ–љ."

    def _extract_text_from_doc(self, file_path: str) -> str:
        """
        –Ш–Ј–≤–ї–µ—З–µ–љ–Є–µ —В–µ–Ї—Б—В–∞ –Є–Ј —Б—В–∞—А–Њ–≥–Њ —Д–Њ—А–Љ–∞—В–∞ .doc.
        –Ш—Б–њ–Њ–ї—М–Ј—Г–µ—В striprtf (–µ—Б–ї–Є —Н—В–Њ RTF) –Є–ї–Є —Б–Є—Б—В–µ–Љ–љ—Г—О –Ї–Њ–Љ–∞–љ–і—Г antiword.
        """
        # 1. –Я–Њ–њ—Л—В–Ї–∞ —З–µ—А–µ–Ј striprtf (–љ–µ–Ї–Њ—В–Њ—А—Л–µ .doc - —Н—В–Њ –љ–∞ —Б–∞–Љ–Њ–Љ –і–µ–ї–µ RTF)
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                if "{\\rtf" in content:
                    text = rtf_to_text(content)
                    if text.strip():
                        logger.info(f"Striprtf extracted {len(text)} characters from .doc (RTF)")
                        return text
        except Exception as e:
            logger.warning(f"Striprtf failed for .doc: {e}")

        # 2. –Я–Њ–њ—Л—В–Ї–∞ —З–µ—А–µ–Ј —Б–Є—Б—В–µ–Љ–љ—Г—О –Ї–Њ–Љ–∞–љ–і—Г antiword
        import shutil
        import subprocess
        
        antiword_cmd = shutil.which('antiword')
        if antiword_cmd:
            try:
                result = subprocess.run([antiword_cmd, file_path], capture_output=True, text=True, errors='ignore')
                if result.returncode == 0 and result.stdout.strip():
                    logger.info(f"Antiword extracted {len(result.stdout)} characters from .doc")
                    return result.stdout
            except Exception as e:
                logger.warning(f"Antiword execution failed: {e}")
        else:
            logger.warning("Antiword executable not found in PATH.")

        # 3. –Я–Њ—Б–ї–µ–і–љ—П—П –њ–Њ–њ—Л—В–Ї–∞: —З—В–µ–љ–Є–µ –Ї–∞–Ї –њ—А–Њ—Б—В–Њ–≥–Њ —В–µ–Ї—Б—В–∞ (–Є–љ–Њ–≥–і–∞ –њ–Њ–Љ–Њ–≥–∞–µ—В –і–ї—П –Њ—З–µ–љ—М —Б—В–∞—А—Л—Е –Є–ї–Є –њ–Њ–≤—А–µ–ґ–і–µ–љ–љ—Л—Е —Д–∞–є–ї–Њ–≤)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                raw_content = f.read()
                # –Я—Л—В–∞–µ–Љ—Б—П –љ–∞–є—В–Є —Е–Њ—В—М –Ї–∞–Ї–Њ–є-—В–Њ –Њ—Б–Љ—Л—Б–ї–µ–љ–љ—Л–є —В–µ–Ї—Б—В —Б—А–µ–і–Є –±–Є–љ–∞—А–љ—Л—Е –і–∞–љ–љ—Л—Е
                import re
                clean_text = re.sub(r'[^\x20-\x7E\u0400-\u04FF\n\t]', ' ', raw_content)
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                if len(clean_text) > 100:
                    logger.info("Extracted partial text from .doc using raw fallback.")
                    return f"[–Т–Э–Ш–Ь–Р–Э–Ш–Х: –Ґ–µ–Ї—Б—В –Є–Ј–≤–ї–µ—З–µ–љ —З–∞—Б—В–Є—З–љ–Њ]\n\n{clean_text}"
        except:
            pass

        msg = f"–Э–µ —Г–і–∞–ї–Њ—Б—М –њ—А–Њ—З–Є—В–∞—В—М —Д–∞–є–ї {os.path.basename(file_path)}."
        if platform.system() == "Windows":
            msg += "\n\n–Ф–Ы–ѓ –Ш–°–Я–†–Р–Т–Ы–Х–Э–Ш–ѓ:\n1. –£—Б—В–∞–љ–Њ–≤–Є—В–µ —Г—В–Є–ї–Є—В—Г Antiword –Є –і–Њ–±–∞–≤—М—В–µ –µ—С –≤ PATH.\n2. –Ш–Ы–Ш (–њ—А–Њ—Й–µ) –њ–µ—А–µ—Б–Њ—Е—А–∞–љ–Є—В–µ —Д–∞–є–ї –≤ —Д–Њ—А–Љ–∞—В–µ .docx."
        else:
            msg += "\n\n–£—Б—В–∞–љ–Њ–≤–Є—В–µ –њ–∞–Ї–µ—В antiword (sudo apt install antiword)."
            
        return f"[–Ю–®–Ш–С–Ъ–Р –§–Ю–†–Ь–Р–Ґ–Р] {msg}"
